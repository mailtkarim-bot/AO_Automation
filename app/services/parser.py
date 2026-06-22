"""Parsing des documents du DCE : PDF, DOCX, XLSX, ZIP (DCE complet).

On privilégie une extraction structurée (page par page quand possible),
pour pouvoir citer une source précise dans les réponses IA.
"""
from __future__ import annotations

import io
import zipfile
from dataclasses import dataclass, field

import pdfplumber
from docx import Document as DocxDocument

from app.services.storage import download_bytes


@dataclass
class ParsedPage:
    page: int
    texte: str


@dataclass
class ParsedDoc:
    nom: str
    pages: list[ParsedPage] = field(default_factory=list)
    texte_brut: str = ""
    nb_pages: int = 0

    def flatten(self) -> str:
        return "\n\n".join(p.texte for p in self.pages if p.texte.strip()) or self.texte_brut


def parse_pdf(data: bytes, nom: str) -> ParsedDoc:
    out = ParsedDoc(nom=nom)
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        out.nb_pages = len(pdf.pages)
        for i, page in enumerate(pdf.pages, start=1):
            texte = page.extract_text() or ""
            # On récupère aussi les tableaux (BPU, DPGF...)
            try:
                tables = page.extract_tables() or []
            except Exception:
                tables = []
            for t in tables:
                for row in t:
                    cells = [c for c in (row or []) if c]
                    if cells:
                        texte += "\n" + " | ".join(cells)
            if texte.strip():
                out.pages.append(ParsedPage(page=i, texte=texte))
    out.texte_brut = out.flatten()
    return out


def parse_docx(data: bytes, nom: str) -> ParsedDoc:
    out = ParsedDoc(nom=nom)
    doc = DocxDocument(io.BytesIO(data))
    chunks: list[str] = []
    # Paragraphes
    for para in doc.paragraphs:
        if para.text.strip():
            chunks.append(para.text)
    # Tableaux (souvent présents en CCTP/BPU)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                chunks.append(" | ".join(cells))
    out.texte_brut = "\n\n".join(chunks)
    out.nb_pages = max(1, len(chunks) // 45)  # estimation grossière
    # Pas de notion de page réelle en DOCX → on regroupe en blocs
    out.pages = [ParsedPage(page=1, texte=out.texte_brut)]
    return out


def parse_xlsx(data: bytes, nom: str) -> ParsedDoc:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    lignes: list[str] = []
    for ws in wb.worksheets:
        lignes.append(f"### Feuille : {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = ["" if c is None else str(c) for c in row]
            if any(c.strip() for c in cells):
                lignes.append(" | ".join(cells))
    texte = "\n".join(lignes)
    return ParsedDoc(nom=nom, pages=[ParsedPage(page=1, texte=texte)],
                     texte_brut=texte, nb_pages=1)


def parse_bytes(data: bytes, nom: str, mime_type: str | None = None) -> ParsedDoc:
    """Dispatcheur principal selon l'extension/mime."""
    nom_lower = nom.lower()
    try:
        if nom_lower.endswith(".pdf") or mime_type == "application/pdf":
            return parse_pdf(data, nom)
        if nom_lower.endswith(".docx") or "wordprocessingml" in (mime_type or ""):
            return parse_docx(data, nom)
        if nom_lower.endswith(".xlsx") or "spreadsheetml" in (mime_type or ""):
            return parse_xlsx(data, nom)
        if nom_lower.endswith(".doc"):
            raise NotImplementedError("Format .doc legacy : convertir en .docx")
        # Fallback texte
        try:
            texte = data.decode("utf-8", errors="ignore")
        except Exception:
            texte = ""
        return ParsedDoc(nom=nom, pages=[ParsedPage(page=1, texte=texte)],
                         texte_brut=texte, nb_pages=1)
    except Exception as e:
        raise RuntimeError(f"Échec du parsing de {nom} : {e}") from e


def unzip_dce(data: bytes) -> list[tuple[str, bytes]]:
    """Un DCE est souvent livré en ZIP. Renvoie (nom_fichier, bytes) pour chaque pièce."""
    fichiers: list[tuple[str, bytes]] = []
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        for info in zf.infolist():
            if info.is_dir():
                continue
            if info.filename.startswith("__MACOSX") or info.filename.endswith(".DS_Store"):
                continue
            fichiers.append((info.filename, zf.read(info.filename)))
    return fichiers
